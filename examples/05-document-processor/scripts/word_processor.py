#!/usr/bin/env python3
"""
Word 文档处理脚本
支持读取、提取文本和获取元数据
"""

import sys
import os

try:
    from docx import Document
except ImportError:
    print("错误: 请安装 python-docx 库")
    print("安装命令: pip install python-docx")
    sys.exit(1)


def read_document(file_path):
    """读取 Word 文档内容"""
    try:
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return {
            "success": True,
            "content": "\n".join(text_content),
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()])
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


def extract_text(file_path):
    """提取文档中的所有文本"""
    result = read_document(file_path)
    if result["success"]:
        return result["content"]
    return None


def get_metadata(file_path):
    """获取文档元数据"""
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        
        return {
            "success": True,
            "title": core_props.title or "未设置",
            "author": core_props.author or "未设置",
            "created": str(core_props.created) if core_props.created else "未设置",
            "modified": str(core_props.modified) if core_props.modified else "未设置"
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("错误: 参数不足")
        print("用法: python word_processor.py <action> <file_path>")
        print("action: read, extract_text, get_metadata")
        sys.exit(1)
    
    action = sys.argv[1]
    file_path = sys.argv[2]
    
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在: {file_path}")
        sys.exit(1)
    
    if action == "read":
        result = read_document(file_path)
        if result["success"]:
            print(f"段落数: {result['paragraph_count']}")
            print("\n内容:")
            print(result["content"])
        else:
            print(f"错误: {result['error']}")
            sys.exit(1)
    
    elif action == "extract_text":
        text = extract_text(file_path)
        if text:
            print(text)
        else:
            print("错误: 无法提取文本")
            sys.exit(1)
    
    elif action == "get_metadata":
        result = get_metadata(file_path)
        if result["success"]:
            print(f"标题: {result['title']}")
            print(f"作者: {result['author']}")
            print(f"创建时间: {result['created']}")
            print(f"修改时间: {result['modified']}")
        else:
            print(f"错误: {result['error']}")
            sys.exit(1)
    
    else:
        print(f"错误: 未知的操作类型: {action}")
        sys.exit(1)


